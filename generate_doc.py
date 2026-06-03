from docx import Document
from docx.shared import Pt
import os

root = os.path.dirname(__file__)
readme_path = os.path.join(root, 'README.md')
proj_root = os.path.join(root, 'Project_65133295')
output_path = os.path.join(root, 'Hệ thống quản lý dãy trọ Thanh Thảo Stay.docx')

def read_file(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception:
        return ''

readme = read_file(readme_path)

# collect models and controllers
models_dir = os.path.join(proj_root, 'Models')
controllers_dir = os.path.join(proj_root, 'Controllers')
models = []
controllers = []
if os.path.isdir(models_dir):
    for fn in sorted(os.listdir(models_dir)):
        if fn.endswith('.cs'):
            models.append(fn)
if os.path.isdir(controllers_dir):
    for fn in sorted(os.listdir(controllers_dir)):
        if fn.endswith('.cs'):
            controllers.append(fn)

# packages (basic list)
packages = []
packages_config = os.path.join(proj_root, 'packages.config')
if os.path.exists(packages_config):
    pc = read_file(packages_config)
    for line in pc.splitlines():
        if 'package id' in line.lower() or 'id="' in line:
            packages.append(line.strip())

# Create document
doc = Document()

# Title
title = doc.add_heading('Hệ thống quản lý dãy trọ Thanh Thảo Stay', level=1)

# Project overview
doc.add_heading('1. Tổng quan dự án', level=2)
if readme:
    doc.add_paragraph(readme[:4000])
else:
    doc.add_paragraph('Dự án là một hệ thống quản lý dãy trọ sử dụng ASP.NET MVC, cung cấp chức năng quản lý phòng, đặt phòng, thanh toán và người dùng.')

# Features
doc.add_heading('2. Yêu cầu & Tính năng', level=2)
doc.add_paragraph('- Quản lý phòng: thêm/sửa/xóa phòng, hình ảnh phòng, trạng thái phòng.')
doc.add_paragraph('- Đặt phòng và thanh toán.')
doc.add_paragraph('- Quản lý người dùng và phân quyền (Admin/User).')

# Architecture
doc.add_heading('3. Kiến trúc & Công nghệ', level=2)
doc.add_paragraph('- Framework: ASP.NET MVC 5')
doc.add_paragraph('- ORM: Entity Framework 6')
doc.add_paragraph('- Frontend: Bootstrap, jQuery')
if packages:
    doc.add_paragraph('Gói NuGet/Packages (tóm tắt):')
    for p in packages[:40]:
        doc.add_paragraph(p, style='List Bullet')

# Database
doc.add_heading('4. Cơ sở dữ liệu', level=2)
doc.add_paragraph('Các model chính (các file trong thư mục Models):')
for m in models:
    doc.add_paragraph('- ' + m, style='List Bullet')

# Controllers
doc.add_heading('5. Luồng chức năng chính (Controllers)', level=2)
doc.add_paragraph('Các controller chính (thư mục Controllers):')
for c in controllers:
    doc.add_paragraph('- ' + c, style='List Bullet')

# UI
doc.add_heading('6. Giao diện người dùng', level=2)
doc.add_paragraph('Giao diện sử dụng Bootstrap. Các view chính nằm trong thư mục Views.')

# Cài đặt & Triển khai
doc.add_heading('7. Cài đặt & Triển khai', level=2)
doc.add_paragraph('1) Mở solution Project_65133295.sln bằng Visual Studio 2019/2022.')
doc.add_paragraph('2) Cài các package NuGet: Restore NuGet packages.')
doc.add_paragraph('3) Cấu hình chuỗi kết nối trong Web.config nếu dùng SQL Server.')

# Kiểm thử
doc.add_heading('8. Kiểm thử', level=2)
doc.add_paragraph('- Kiểm thử chức năng đăng nhập/đăng ký, đặt phòng, thanh toán.')

# Bảo trì
doc.add_heading('9. Bảo trì & Mở rộng', level=2)
doc.add_paragraph('- Đề xuất: tách service, thêm API REST, CI/CD.')

# Phụ lục
doc.add_heading('Phụ lục', level=2)
doc.add_paragraph('Danh sách file tham khảo:')
if os.path.isdir(proj_root):
    for rootdir, dirs, files in os.walk(proj_root):
        for f in files:
            if f.endswith('.cs') or f.endswith('.config') or f.endswith('.sql'):
                rel = os.path.relpath(os.path.join(rootdir,f), start=os.path.dirname(__file__))
                doc.add_paragraph('- ' + rel, style='List Bullet')

# Save
try:
    doc.save(output_path)
    print('SAVED:'+output_path)
except Exception as e:
    print('ERROR:'+str(e))
